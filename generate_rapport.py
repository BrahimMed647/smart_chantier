#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Rapport de Fin d'Etudes - Smart Chantier
Format identique au rapport de reference (Bouchra Cheikh sidi.pdf)
3 chapitres
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION

# ── Couleurs ─────────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x00, 0x32, 0x66)
BLUE       = RGBColor(0x1F, 0x49, 0x7D)
GRAY       = RGBColor(0x59, 0x59, 0x59)
BLACK      = RGBColor(0x00, 0x00, 0x00)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE = RGBColor(0xBD, 0xD7, 0xEE)
TABLE_HDR  = RGBColor(0x1F, 0x49, 0x7D)

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    tcPr.append(shd)

def set_table_borders(table, color='B8B8B8', sz=4):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    tb = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        e = OxmlElement(f'w:{side}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        tb.append(e)
    tblPr.append(tb)

def cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        e = OxmlElement(f'w:{side}')
        e.set(qn('w:w'), str(val)); e.set(qn('w:type'), 'dxa')
        tcMar.append(e)
    tcPr.append(tcMar)

def add_page_number(footer_para):
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run.font.size = Pt(10)

def set_para_spacing(para, before=0, after=6, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line)

def para_text(doc, text, size=11, bold=False, italic=False, color=None,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=6, indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    set_para_spacing(p, before, after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    else: r.font.color.rgb = BLACK
    return p

def chapter_title(doc, text):
    """Grand titre de chapitre avec ligne de separation"""
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=40, after=10)
    r = p.add_run(text)
    r.font.size = Pt(18); r.bold = True; r.font.color.rgb = DARK_BLUE
    # Ligne sous le titre
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '12')
    bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), '003366')
    pBdr.append(bot); pPr.append(pBdr)
    doc.add_paragraph()  # espace apres

def section_title(doc, text, level=1):
    """Titre de section numerotee"""
    p = doc.add_paragraph()
    set_para_spacing(p, before=14, after=4)
    r = p.add_run(text)
    if level == 1:
        r.font.size = Pt(13); r.bold = True; r.font.color.rgb = DARK_BLUE
    elif level == 2:
        r.font.size = Pt(12); r.bold = True; r.font.color.rgb = BLUE
    else:
        r.font.size = Pt(11); r.bold = True; r.font.color.rgb = GRAY
    return p

def conclusion_intro(doc, text):
    """Paragraphe de conclusion/introduction de section en italique"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, before=8, after=8)
    r = p.add_run(text)
    r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GRAY
    return p

def figure_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=4, after=12)
    r = p.add_run(text)
    r.italic = True; r.bold = True; r.font.size = Pt(10); r.font.color.rgb = DARK_BLUE
    return p

def table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=12, after=4)
    r = p.add_run(text)
    r.italic = True; r.bold = True; r.font.size = Pt(10); r.font.color.rgb = DARK_BLUE
    return p

def bullet_item(doc, text, indent=1.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, before=0, after=3)
    p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(f"- {text}")
    r.font.size = Pt(11); r.font.color.rgb = BLACK
    return p

def make_table(doc, headers, rows, col_widths=None, hdr_color='1F497D'):
    """Cree un tableau avec en-tete colore"""
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t, 'AAAAAA', 6)
    # En-tete
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_bg(c, hdr_color)
        cell_margins(c)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(11)
    # Donnees
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            cell_margins(c)
            if ri % 2 == 1:
                set_cell_bg(c, 'F2F2F2')
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val)); r.font.size = Pt(10); r.font.color.rgb = BLACK
    # Largeurs colonnes
    if col_widths:
        for row in t.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[i])
    doc.add_paragraph()
    return t

def screenshot_box(doc, title, fig_num):
    """Boite placeholder pour capture d'ecran"""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t, 'CCCCCC', 4)
    c = t.rows[0].cells[0]
    set_cell_bg(c, 'F5F5F5')
    c.height = Cm(8)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(p, before=30, after=30)
    r = p.add_run(f"\n[ Capture d'ecran : {title} ]\n")
    r.font.size = Pt(11); r.italic = True; r.font.color.rgb = GRAY
    figure_caption(doc, f"Figure {fig_num} : {title}")

# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()

# Format A4
sec = doc.sections[0]
sec.page_width  = Cm(21)
sec.page_height = Cm(29.7)
sec.left_margin = sec.right_margin = Cm(2.5)
sec.top_margin  = Cm(2.5)
sec.bottom_margin = Cm(2.5)

# Style par defaut
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# Footer avec numero de page
for section in doc.sections:
    footer = section.footer
    if footer.paragraphs:
        fp = footer.paragraphs[0]
    else:
        fp = footer.add_paragraph()
    add_page_number(fp)

# ═══════════════════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ═══════════════════════════════════════════════════════════════════════════════
# En-tete universite
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Ministere de l'Enseignement Superieur et de la Recherche Scientifique")
r.font.size = Pt(11); r.bold = True; r.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Universite / Ecole Superieure d'Informatique")
r.font.size = Pt(12); r.bold = True; r.font.color.rgb = DARK_BLUE

doc.add_paragraph()

# Annee universitaire
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Annee universitaire\n2024 - 2025")
r.font.size = Pt(11); r.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_paragraph()

# Type de rapport
t_cover = doc.add_table(rows=1, cols=1)
t_cover.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t_cover, '003366', 10)
c = t_cover.rows[0].cells[0]
set_cell_bg(c, 'E8EFF6')
cell_margins(c, 200, 200, 300, 300)
p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Rapport de stage de fin d'etudes pour l'obtention\ndu diplome de Master en Genie Logiciel")
r.font.size = Pt(13); r.bold = True; r.font.color.rgb = DARK_BLUE

doc.add_paragraph()

# Titre du projet
t2 = doc.add_table(rows=1, cols=1)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t2, '003366', 14)
c2 = t2.rows[0].cells[0]
set_cell_bg(c2, '003366')
cell_margins(c2, 200, 200, 300, 300)
p = c2.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Theme :\nConception et developpement d'une application\nmobile et web de gestion de chantiers de\nconstruction (Smart Chantier)")
r.font.size = Pt(14); r.bold = True; r.font.color.rgb = WHITE

doc.add_paragraph()
doc.add_paragraph()

# Auteur et encadreurs
t3 = doc.add_table(rows=3, cols=2)
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t3, 'CCCCCC', 4)
infos = [
    ("Elabore par", "Brahim HAMOUD"),
    ("Encadrant universitaire :", "Nom de l'encadreur"),
    ("Encadrant professionnel :", "Nom de l'encadreur professionnel"),
]
for i, (lbl, val) in enumerate(infos):
    c0 = t3.rows[i].cells[0]; c1 = t3.rows[i].cells[1]
    cell_margins(c0); cell_margins(c1)
    set_cell_bg(c0, 'E8EFF6')
    c0.paragraphs[0].add_run(lbl).bold = True
    c0.paragraphs[0].runs[0].font.size = Pt(11)
    c0.paragraphs[0].runs[0].font.color.rgb = DARK_BLUE
    c1.paragraphs[0].add_run(val)
    c1.paragraphs[0].runs[0].font.size = Pt(11)
    c1.paragraphs[0].runs[0].font.color.rgb = BLACK

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# DEDICACE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p, before=0, after=20)
r = p.add_run("Dedicace")
r.font.size = Pt(18); r.bold = True; r.font.color.rgb = DARK_BLUE

# Ligne sous le titre
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot = OxmlElement('w:bottom')
bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '12')
bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), '003366')
pBdr.append(bot); pPr.append(pBdr)

doc.add_paragraph()

for para in [
    "Je dedie ce rapport de stage a mes parents bien-aimes, piliers de ma vie et sources inepuisables d'inspiration.",
    "A ma chere et tendre maman,",
    "Ta douceur, ton amour inconditionnel et tes prieres silencieuses m'ont accompagne a chaque etape de mon parcours. Tu as su, par ta tendresse et ta force interieure, m'apporter le courage necessaire pour avancer, meme dans les moments de doute. Merci pour tes sacrifices quotidiens et ton soutien sans faille.",
    "A mon cher papa,",
    "Ta sagesse, ton calme, et ta confiance en moi ont ete pour moi un phare dans les moments incertains. Tu m'as transmis des valeurs precieuses : la rigueur, l'honnetete, et surtout la perseverance. Par ton exemple, tu m'as appris que rien ne remplace le travail bien fait et l'effort constant.",
    "Ce travail n'est qu'un modeste reflet de tout ce que vous m'avez donne.",
    "Merci pour votre patience, vos prieres, vos encouragements, et votre foi inebranlable en moi. Que Dieu vous protege, vous comble de Sa misericorde, et vous accorde une longue vie pleine de sante et de paix.",
]:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, before=0, after=8)
    p.runs[0].font.size = Pt(12)
    if para in ["A ma chere et tendre maman,", "A mon cher papa,"]:
        p.runs[0].bold = True; p.runs[0].font.color.rgb = DARK_BLUE

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# REMERCIEMENTS
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p, before=0, after=20)
r = p.add_run("REMERCIEMENTS")
r.font.size = Pt(18); r.bold = True; r.font.color.rgb = DARK_BLUE
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot = OxmlElement('w:bottom'); bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '12')
bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), '003366')
pBdr.append(bot); pPr.append(pBdr)
doc.add_paragraph()

for para in [
    "Je tiens a exprimer ma sincere gratitude a mon encadreur academique pour son accompagnement precieux, sa disponibilite et ses conseils tout au long de la realisation de ce projet. Son expertise et son soutien constant ont ete determinants dans la reussite de ce travail et dans mon developpement personnel et professionnel.",
    "Je souhaite egalement remercier chaleureusement l'ensemble du corps enseignant du departement informatique pour la qualite de la formation dispensee tout au long de ces annees d'etudes. Leurs enseignements m'ont fourni les bases theoriques et pratiques indispensables a la realisation de ce projet.",
    "Mes remerciements s'adressent egalement aux membres du jury pour avoir accepte d'evaluer ce travail et pour le temps precieux qu'ils lui ont consacre. Leurs remarques et suggestions contribueront sans nul doute a l'amelioration de ce travail.",
    "Enfin, je tiens a remercier toutes les personnes qui ont, de pres ou de loin, participe a l'aboutissement de ce projet. Votre confiance, votre soutien et vos encouragements ont enrichi cette experience et m'ont permis de grandir en tant que futur professionnel.",
]:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, before=0, after=10)
    p.runs[0].font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# TABLE DES MATIERES
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p, before=0, after=20)
r = p.add_run("Table des matieres")
r.font.size = Pt(18); r.bold = True; r.font.color.rgb = DARK_BLUE
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot = OxmlElement('w:bottom'); bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '12')
bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), '003366')
pBdr.append(bot); pPr.append(pBdr)
doc.add_paragraph()

toc_entries = [
    ("Liste de figures",                                             False, 0),
    ("Liste de tableaux",                                           False, 0),
    ("Liste des abreviations",                                      False, 0),
    ("Introduction Generale",                                       False, 0),
    ("Chapitre 1 : Presentation du Projet",                         True,  0),
    ("1.1 Introduction",                                            False, 1),
    ("1.2 Description du projet",                                   False, 1),
    ("1.3 Problematique",                                           False, 1),
    ("1.4 Objectifs",                                               False, 1),
    ("1.5 Conclusion",                                              False, 1),
    ("Chapitre 2 : Analyse des Besoins et Conception",              True,  0),
    ("2.1 Introduction",                                            False, 1),
    ("2.2 Analyse des besoins",                                     False, 1),
    ("2.2.1 Besoins fonctionnels",                                  False, 2),
    ("2.2.2 Besoins non fonctionnels",                              False, 2),
    ("2.3 Modelisation UML",                                        False, 1),
    ("2.3.1 Identification des acteurs",                            False, 2),
    ("2.3.2 Diagramme de cas d'utilisation",                        False, 2),
    ("2.3.3 Diagramme de classes",                                  False, 2),
    ("2.3.4 Diagramme de sequence - Authentification",              False, 2),
    ("2.4 Architecture du systeme",                                 False, 1),
    ("2.4.1 Architecture globale",                                  False, 2),
    ("2.4.2 Architecture Flutter (Pattern Provider)",               False, 2),
    ("2.4.3 Architecture Django REST Framework",                    False, 2),
    ("2.4.4 Modele de base de donnees (ERD)",                       False, 2),
    ("2.5 Conclusion",                                              False, 1),
    ("Chapitre 3 : Realisation et Implementation",                  True,  0),
    ("3.1 Introduction",                                            False, 1),
    ("3.2 Environnement de travail",                                False, 1),
    ("3.2.1 Outils de conception",                                  False, 2),
    ("3.2.2 Outils de developpement",                               False, 2),
    ("3.3 Presentation des interfaces",                             False, 1),
    ("3.3.1 Interface de connexion",                                False, 2),
    ("3.3.2 Tableau de bord",                                       False, 2),
    ("3.3.3 Gestion des projets",                                   False, 2),
    ("3.3.4 Gestion des taches",                                    False, 2),
    ("3.3.5 Rapports journaliers",                                  False, 2),
    ("3.3.6 Gestion des depenses",                                  False, 2),
    ("3.3.7 Systeme d'alertes",                                     False, 2),
    ("3.4 Tests et validation",                                     False, 1),
    ("3.5 Les apports du projet",                                   False, 1),
    ("3.6 Conclusion",                                              False, 1),
    ("Conclusion Generale",                                         False, 0),
    ("Bibliographie",                                               False, 0),
    ("Annexes",                                                     False, 0),
]

for entry, is_chapter, indent_level in toc_entries:
    p = doc.add_paragraph()
    set_para_spacing(p, before=0, after=2)
    p.paragraph_format.left_indent = Cm(indent_level * 0.7)
    r = p.add_run(entry)
    r.font.size = Pt(12 if is_chapter else 11)
    r.bold = is_chapter
    r.font.color.rgb = DARK_BLUE if is_chapter else BLACK

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# LISTE DES FIGURES
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p, before=0, after=20)
r = p.add_run("Liste de figures")
r.font.size = Pt(16); r.bold = True; r.font.color.rgb = DARK_BLUE
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot = OxmlElement('w:bottom'); bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '8')
bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), '003366')
pBdr.append(bot); pPr.append(pBdr)

figures = [
    "Figure 1 : Architecture globale du systeme Smart Chantier (3-Tiers)",
    "Figure 2 : Diagramme de cas d'utilisation",
    "Figure 3 : Diagramme de classes",
    "Figure 4 : Diagramme de sequence - Authentification JWT",
    "Figure 5 : Modele Entite-Association (ERD)",
    "Figure 6 : Architecture Flutter - Pattern Provider",
    "Figure 7 : Architecture Django REST Framework",
    "Figure 8 : Logo UML",
    "Figure 9 : Logo Draw.io",
    "Figure 10 : Logo Flutter",
    "Figure 11 : Logo Dart",
    "Figure 12 : Logo Visual Studio Code",
    "Figure 13 : Logo Django",
    "Figure 14 : Interface de connexion",
    "Figure 15 : Tableau de bord principal",
    "Figure 16 : Gestion des projets - liste",
    "Figure 17 : Detail d'un projet",
    "Figure 18 : Gestion des taches",
    "Figure 19 : Rapports journaliers",
    "Figure 20 : Gestion des depenses",
    "Figure 21 : Systeme d'alertes",
]
for f in figures:
    p = doc.add_paragraph(f)
    p.runs[0].font.size = Pt(11)
    set_para_spacing(p, before=0, after=2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# LISTE DES TABLEAUX
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p, before=0, after=20)
r = p.add_run("Liste de tableaux")
r.font.size = Pt(16); r.bold = True; r.font.color.rgb = DARK_BLUE
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot = OxmlElement('w:bottom'); bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '8')
bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), '003366')
pBdr.append(bot); pPr.append(pBdr)

tableaux = [
    "Tableau 1 : Acteurs du systeme et leurs roles",
    "Tableau 2 : Comparaison des solutions existantes",
    "Tableau 3 : Besoins non fonctionnels",
    "Tableau 4 : Technologies et outils utilises",
    "Tableau 5 : Resultats des tests fonctionnels",
    "Tableau 6 : Endpoints API - Liste complete",
]
for t in tableaux:
    p = doc.add_paragraph(t)
    p.runs[0].font.size = Pt(11)
    set_para_spacing(p, before=0, after=2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# LISTE DES ABREVIATIONS
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p, before=0, after=20)
r = p.add_run("Liste des abreviations")
r.font.size = Pt(16); r.bold = True; r.font.color.rgb = DARK_BLUE
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot = OxmlElement('w:bottom'); bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '8')
bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), '003366')
pBdr.append(bot); pPr.append(pBdr)
doc.add_paragraph()

abrevs = [
    ("API",    "Application Programming Interface"),
    ("BTP",    "Batiment et Travaux Publics"),
    ("CRUD",   "Create, Read, Update, Delete"),
    ("DRF",    "Django REST Framework"),
    ("ERD",    "Entity-Relationship Diagram"),
    ("IDE",    "Integrated Development Environment"),
    ("JSON",   "JavaScript Object Notation"),
    ("JWT",    "JSON Web Token"),
    ("MVC",    "Model-View-Controller"),
    ("MVVM",   "Model-View-ViewModel"),
    ("ORM",    "Object-Relational Mapping"),
    ("REST",   "Representational State Transfer"),
    ("SDK",    "Software Development Kit"),
    ("UI",     "User Interface"),
    ("UML",    "Unified Modeling Language"),
    ("URL",    "Uniform Resource Locator"),
    ("UX",     "User Experience"),
]
make_table(doc, ["Abreviation", "Signification"], abrevs, col_widths=[4, 12])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# INTRODUCTION GENERALE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p, before=0, after=20)
r = p.add_run("Introduction Generale")
r.font.size = Pt(18); r.bold = True; r.font.color.rgb = DARK_BLUE
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot = OxmlElement('w:bottom'); bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '12')
bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), '003366')
pBdr.append(bot); pPr.append(pBdr)
doc.add_paragraph()

for para in [
    "Dans le cadre de la formation en Genie Logiciel, un projet de fin d'etudes est prevu en derniere annee. Ce projet constitue une opportunite essentielle pour appliquer les connaissances theoriques acquises tout au long de la formation, et pour concevoir une solution logicielle repondant a un besoin reel.",
    "Le secteur du batiment et des travaux publics (BTP) est l'un des domaines les plus complexes a gerer en raison du nombre eleve d'intervenants, de la multiplicite des taches a coordonner, et des contraintes budgetaires et temporelles strictes. La gestion manuelle des chantiers, encore tres repandue, engendre souvent des retards, des depassements de budget et une communication insuffisante entre les parties prenantes.",
    "Face a ces defis, nous avons choisi de developper Smart Chantier, une application mobile et web complete dediee a la gestion intelligente des chantiers de construction. Cette solution permet aux chefs de projet, ingenieurs et administrateurs de suivre l'avancement des travaux, controler les budgets, gerer les equipes, documenter les incidents et generer des rapports automatiquement depuis leur smartphone ou navigateur web.",
    "Cette experience nous a permis de consolider nos competences techniques en programmation mobile (Flutter/Dart), en developpement back-end (Python/Django), en conception d'interfaces (UX/UI), ainsi que nos aptitudes a structurer un projet digital de A a Z.",
    "Ce rapport rend compte des differentes etapes du projet, des technologies utilisees, des interfaces realisees, et des resultats obtenus. Il est organise comme suit :",
]:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, before=0, after=10)
    p.runs[0].font.size = Pt(12)

for item in [
    "Chapitre 1 : Presentation du projet (contexte, description, problematique, objectifs)",
    "Chapitre 2 : Analyse des besoins et conception (UML, architecture, modelisation)",
    "Chapitre 3 : Realisation et implementation (environnement, interfaces, tests)",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_para_spacing(p, before=0, after=4)
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(f"- {item}")
    r.font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 1 : PRESENTATION DU PROJET
# ═══════════════════════════════════════════════════════════════════════════════
chapter_title(doc, "Chapitre 1 : Presentation du Projet")

section_title(doc, "1.1 Introduction", 1)
p = doc.add_paragraph(
    "Ce chapitre a pour but de presenter le projet Smart Chantier. Il s'agit d'un projet "
    "de developpement d'une application mobile et web destinee a faciliter la gestion des "
    "chantiers de construction. Cette solution numerique vise a repondre aux limites des "
    "methodes traditionnelles encore largement utilisees dans le secteur. La description du "
    "projet, la problematique identifiee, ainsi que les objectifs poursuivis seront detailles "
    "afin de mieux comprendre l'interet et l'impact de cette application."
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=10)
p.runs[0].font.size = Pt(12)

section_title(doc, "1.2 Description du Projet", 1)
for para in [
    "Le projet Smart Chantier s'inscrit dans une volonte d'apporter une solution numerique aux problemes rencontres dans la gestion des chantiers de construction. En effet, de nombreuses entreprises du BTP continuent a utiliser des methodes traditionnelles comme les tableurs Excel, les carnets papier ou les appels telephoniques pour gerer leurs projets, ce qui entraine des erreurs, des oublis, des conflits de planning et une perte de temps considerable.",
    "Face a ces difficultes, Smart Chantier propose une application mobile moderne, destinee a faciliter et automatiser la gestion des chantiers aussi bien pour les administrateurs que pour les chefs de projet, ingenieurs et techniciens sur le terrain. Cette application repond a un besoin reel d'optimisation du service, de reduction des erreurs humaines et de gain de temps, tout en ameliorant la communication et la traçabilite des actions.",
    "L'application est developpee selon une architecture client-serveur moderne : un front-end Flutter (framework multi-plateforme de Google) communiquant via une API RESTful developpee avec Django REST Framework (Python). L'authentification est securisee par des tokens JWT, et les donnees sont stockees dans une base SQLite en developpement et PostgreSQL en production.",
    "Le projet represente ainsi une etape importante dans la digitalisation du secteur de la construction, et s'inscrit dans une logique d'innovation et d'amelioration continue des processus de gestion.",
]:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=10)
    p.runs[0].font.size = Pt(12)

section_title(doc, "1.3 Problematique", 1)
for para in [
    "Meme si les nouvelles technologies sont tres utilisees aujourd'hui, beaucoup d'entreprises du BTP gerent encore leurs chantiers de facon traditionnelle, en utilisant des tableurs, des carnets papier ou des appels telephoniques. Cette methode pose plusieurs problemes :",
]:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=8)
    p.runs[0].font.size = Pt(12)

for item, detail in [
    ("Lenteur et inefficacite :", "Gerer les informations manuellement prend du temps et devient difficile quand plusieurs projets se deroulent en parallele, ce qui peut engendrer des retards et des pertes de clients."),
    ("Erreurs frequentes :", "Il arrive souvent qu'on se trompe en notant un budget, une date d'echeance ou l'avancement d'une tache, ce qui peut creer des malentendus et des conflits."),
    ("Manque de visibilite :", "Sans outil numerique, il est difficile d'avoir une vue globale en temps reel sur l'etat de tous les chantiers, les depenses engagees ou les alertes a traiter."),
    ("Suivi difficile :", "Modifier les informations d'un projet, affecter une tache ou envoyer une alerte peut etre complique sans application dediee. Cela demande beaucoup de travail pour les responsables."),
    ("Perte d'informations :", "Les documents papier peuvent etre perdus, deteriores ou oublies, ce qui compromet la tracabilite des actions et des decisions prises sur les chantiers."),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, before=0, after=5)
    p.paragraph_format.left_indent = Cm(1)
    r1 = p.add_run(item); r1.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = DARK_BLUE
    r2 = p.add_run(f" {detail}"); r2.font.size = Pt(12)

p = doc.add_paragraph(
    "Pour toutes ces raisons, developper une application mobile et web qui aide a gerer les "
    "chantiers devient necessaire. Elle permet de gagner du temps, d'eviter les erreurs, de "
    "mieux organiser les projets, et d'ameliorer la communication entre toutes les parties prenantes."
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, before=8, after=10)
p.runs[0].font.size = Pt(12)

section_title(doc, "1.4 Objectifs", 1)
p = doc.add_paragraph("Les objectifs principaux de ce projet sont les suivants :")
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=5)
p.runs[0].font.size = Pt(12)

for obj in [
    "Permettre le suivi en temps reel de l'avancement de plusieurs projets simultanement.",
    "Centraliser la gestion des taches pour eviter les erreurs, les doublons et assurer un bon suivi.",
    "Controle budgetaire rigoureux avec alertes automatiques en cas de depassement.",
    "Generer des rapports journaliers automatiquement pour documenter l'avancement des travaux.",
    "Centraliser le systeme d'alertes pour signaler tout incident, retard ou probleme sur le chantier.",
    "Permettre la documentation photographique des chantiers avec geolocalization.",
    "Assurer l'acces multi-plateforme : smartphone Android/iOS et navigateur web.",
    "Garantir la securite des donnees : authentification JWT, gestion des roles et des permissions.",
]:
    bullet_item(doc, obj)

section_title(doc, "1.5 Conclusion", 1)
conclusion_intro(doc,
    "En resume, ce projet s'inscrit dans une demarche de modernisation du secteur de la "
    "construction, en proposant une application mobile et web innovante pour gerer les chantiers. "
    "Il repond a un besoin concret exprime par les professionnels du BTP, en apportant des solutions "
    "aux problemes de lenteur, d'erreurs et de manque de visibilite. Les objectifs fixes montrent "
    "clairement l'ambition d'ameliorer l'efficacite, la communication et l'experience utilisateur. "
    "Cette base solide permettra de passer a l'analyse des besoins et a la conception technique "
    "dans les prochaines etapes du developpement."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 2 : ANALYSE DES BESOINS ET CONCEPTION
# ═══════════════════════════════════════════════════════════════════════════════
chapter_title(doc, "Chapitre 2 : Analyse des Besoins et Conception")

section_title(doc, "2.1 Introduction", 1)
p = doc.add_paragraph(
    "Ce chapitre a pour objectif de presenter les etapes d'analyse des besoins ainsi que "
    "la conception du systeme Smart Chantier. Il s'agit d'identifier les besoins fonctionnels "
    "et non fonctionnels des utilisateurs, puis de les traduire en modeles visuels a l'aide "
    "de diagrammes UML adaptes. Cette etape permet de structurer le fonctionnement de "
    "l'application avant son developpement, en precisantles interactions entre les acteurs "
    "et le systeme, ainsi que l'organisation des donnees."
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=10)
p.runs[0].font.size = Pt(12)

# ── 2.2 Analyse des besoins ──
section_title(doc, "2.2 Analyse des besoins", 1)
section_title(doc, "2.2.1 Besoins fonctionnels", 2)
p = doc.add_paragraph(
    "Ce sont les fonctionnalites principales que l'application doit offrir pour repondre "
    "aux attentes des utilisateurs :"
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=6)
p.runs[0].font.size = Pt(12)

for num, func, items in [
    ("1.", "Authentification et gestion des utilisateurs", [
        "Connexion securisee par email et mot de passe avec tokens JWT.",
        "Gestion des roles : Administrateur, Chef de projet, Ingenieur, Technicien.",
        "Modification du profil : nom, email, telephone, mot de passe.",
        "Auto-connexion si un token valide est deja stocke.",
    ]),
    ("2.", "Gestion des projets", [
        "Creer, modifier et consulter des projets de construction.",
        "Suivi du pourcentage d'avancement et des statuts (En attente, En cours, Termine, Annule).",
        "Calcul automatique du budget restant et detection des depassements.",
        "Visualisation de la localisation GPS du chantier.",
    ]),
    ("3.", "Gestion des taches", [
        "Creer des taches liees a un projet avec titre, description et date d'echeance.",
        "Affecter les taches aux membres de l'equipe.",
        "Suivre la priorite : Faible, Normale, Haute, Urgente.",
        "Mettre a jour le statut et le pourcentage de realisation.",
    ]),
    ("4.", "Rapports journaliers", [
        "Creer des rapports quotidiens avec description des travaux realises.",
        "Enregistrer les conditions meteo, l'effectif present et les incidents.",
        "Attacher des photos au rapport.",
        "Filtrer les rapports par projet et par periode.",
    ]),
    ("5.", "Gestion des depenses", [
        "Enregistrer les depenses par categorie : Materiaux, Main d'oeuvre, Equipement, Transport.",
        "Suivre le budget consomme en temps reel par projet.",
        "Visualiser la repartition des depenses par categorie.",
    ]),
    ("6.", "Systeme d'alertes", [
        "Creer des alertes automatiques en cas de depassement de budget ou de retard.",
        "Niveaux de criticite : Information, Avertissement, Danger, Critique.",
        "Resoudre et archiver les alertes traitees.",
        "Badge counter affichant le nombre d'alertes non lues.",
    ]),
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, before=8, after=2)
    r1 = p.add_run(f"{num} {func}"); r1.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = DARK_BLUE
    for item in items:
        bullet_item(doc, item, indent=1.5)

section_title(doc, "2.2.2 Besoins non fonctionnels", 2)
p = doc.add_paragraph(
    "En plus des fonctionnalites principales, l'application doit respecter plusieurs criteres "
    "de qualite afin d'assurer une bonne experience utilisateur :"
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=6)
p.runs[0].font.size = Pt(12)

table_caption(doc, "Tableau 3 : Besoins non fonctionnels")
make_table(doc, ["Critere", "Exigence"], [
    ("Performance",    "Temps de reponse API inferieur a 500ms, interface fluide"),
    ("Securite",       "Chiffrement JWT, stockage securise des tokens, gestion des permissions"),
    ("Fiabilite",      "Disponibilite permanente, donnees accessibles hors-ligne"),
    ("Ergonomie",      "Interface intuitive, navigation claire, adaptation mobile et desktop"),
    ("Portabilite",    "Compatible Android, iOS, Chrome, Firefox, Microsoft Edge"),
    ("Maintenabilite", "Code modulaire, architecture en couches, documentation complete"),
    ("Scalabilite",    "Architecture extensible pour supporter plusieurs organisations"),
], col_widths=[5, 11])

# ── 2.3 Modelisation UML ──
section_title(doc, "2.3 Modelisation UML", 1)
p = doc.add_paragraph(
    "Dans le cadre de l'analyse du projet avant la phase de developpement, quatre types "
    "de diagrammes UML ont ete realises : le diagramme de cas d'utilisation, le diagramme "
    "de classes, le diagramme de sequence. Ces outils permettent de mieux comprendre "
    "les interactions entre les acteurs et le systeme, ainsi que la structure logique des donnees."
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=10)
p.runs[0].font.size = Pt(12)

section_title(doc, "2.3.1 Identification des acteurs", 2)
table_caption(doc, "Tableau 1 : Acteurs du systeme et leurs roles")
make_table(doc, ["Acteur", "Role"], [
    ("Administrateur", "Acces complet : gestion des utilisateurs, organisations, projets, rapports, alertes et budgets"),
    ("Chef de projet",  "Creation et gestion de projets, validation des rapports, controle du budget, affectation des equipes"),
    ("Ingenieur",       "Mise a jour des taches, creation de rapports journaliers, saisie des depenses, upload de photos"),
    ("Technicien",      "Consultation et mise a jour de ses taches affectees, lecture des informations du projet"),
], col_widths=[4, 12])

section_title(doc, "2.3.2 Diagramme de cas d'utilisation", 2)
screenshot_box(doc, "Diagramme de cas d'utilisation", 2)
p = doc.add_paragraph(
    "Commentaire : Ce diagramme presente les differentes interactions possibles entre les "
    "acteurs et l'application Smart Chantier. L'administrateur dispose d'un acces complet "
    "a toutes les fonctionnalites. Le chef de projet peut gerer les projets et valider les "
    "rapports. L'ingenieur peut creer des rapports et saisir des depenses. Le technicien "
    "peut uniquement consulter et mettre a jour ses taches. Certaines actions necessitent "
    "une authentification prealable."
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=10)
p.runs[0].font.size = Pt(12); p.runs[0].italic = True

section_title(doc, "2.3.3 Diagramme de classes", 2)
screenshot_box(doc, "Diagramme de classes", 3)
p = doc.add_paragraph(
    "Commentaire : Ce diagramme represente la structure principale de l'application. "
    "On y trouve les classes User, Organization, Project, Task, DailyReport, Expense, "
    "Alert et SitePhoto, ainsi que leurs relations. Un utilisateur appartient a une "
    "organisation et peut gerer plusieurs projets. Un projet possede des taches, des "
    "rapports journaliers, des depenses, des alertes et des photos. Ce modele assure "
    "une bonne organisation des donnees pour l'application."
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=10)
p.runs[0].font.size = Pt(12); p.runs[0].italic = True

section_title(doc, "2.3.4 Diagramme de sequence - Authentification", 2)
p = doc.add_paragraph(
    "Le diagramme de sequence suivant illustre le flux complet d'authentification entre "
    "l'application Flutter et le backend Django avec les tokens JWT :"
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=6)
p.runs[0].font.size = Pt(12)

seq_data = [
    ["Flutter App",              "AuthProvider",                       "API Django",                    "Base de donnees"],
    ["Saisit email/mot de passe","-> login(email, password)",          "",                              ""],
    ["",                         "POST /api/auth/login/",              "-->",                           ""],
    ["",                         "",                                   "--> SELECT User WHERE email",   ""],
    ["",                         "",                                   "<-- Utilisateur trouve",        ""],
    ["",                         "",                                   "Genere access + refresh token", ""],
    ["",                         "<-- {access, refresh, user_data}",   "",                              ""],
    ["",                         "Stockage token (SecureStorage)",     "",                              ""],
    ["<-- Navigation Dashboard", "",                                    "",                              ""],
    ["",                         "GET /api/projects/ [Bearer Token]",  "-->",                           ""],
    ["",                         "",                                   "Verification JWT valide",       ""],
    ["",                         "<-- Liste des projets (JSON)",        "",                              ""],
    ["Affichage tableau de bord","",                                    "",                              ""],
]
t_seq = doc.add_table(rows=len(seq_data), cols=4)
t_seq.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t_seq, 'AAAAAA', 4)
hdr_colors_seq = ['1F497D', '003366', 'FF6D00', '607D8B']
for ci, (h, col) in enumerate(zip(seq_data[0], hdr_colors_seq)):
    c = t_seq.rows[0].cells[ci]
    set_cell_bg(c, col); cell_margins(c, 60, 60, 80, 80)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(9)
for ri, step in enumerate(seq_data[1:], 1):
    for ci, val in enumerate(step):
        c = t_seq.rows[ri].cells[ci]
        if ri % 2 == 0: set_cell_bg(c, 'F2F2F2')
        cell_margins(c, 40, 40, 60, 60)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(val); r.font.size = Pt(8)
figure_caption(doc, "Figure 4 : Diagramme de sequence - Authentification JWT")

# ── 2.4 Architecture ──
section_title(doc, "2.4 Architecture du systeme", 1)
section_title(doc, "2.4.1 Architecture globale", 2)
p = doc.add_paragraph(
    "Smart Chantier repose sur une architecture client-serveur a trois niveaux (3-Tier Architecture) : "
    "la couche presentation (Flutter), la couche metier (Django REST Framework) et la couche de "
    "donnees (SQLite / PostgreSQL). Cette separation claire des responsabilites facilite la "
    "maintenance et l'evolution de l'application."
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=10)
p.runs[0].font.size = Pt(12)

t_arch = doc.add_table(rows=5, cols=1)
t_arch.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(t_arch, '1F497D', 8)
arch_layers = [
    ("COUCHE PRESENTATION",
     "Application Flutter Mobile (Android / iOS)\nApplication Flutter Web (Chrome, Edge, Firefox)",
     "1F497D", "E8EFF6"),
    ("", "  HTTP / HTTPS  |  API REST  |  Format JSON  |  Token JWT Bearer  ", "FFFFFF", "FFFFFF"),
    ("COUCHE METIER (BACK-END)",
     "Django REST Framework  |  Authentification JWT\nSerialization  |  Permissions par role  |  CORS",
     "003366", "EBF0FA"),
    ("", "  Django ORM  |  Requetes SQL  ", "FFFFFF", "FFFFFF"),
    ("COUCHE DONNEES",
     "SQLite (Developpement)  |  PostgreSQL (Production)\nMedias : Photos de chantier",
     "FF6D00", "FFF3E0"),
]
for i, (title, content, txt_col, bg) in enumerate(arch_layers):
    c = t_arch.rows[i].cells[0]; set_cell_bg(c, bg); cell_margins(c, 120, 120, 200, 200)
    if title:
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title + "\n"); r.bold = True; r.font.size = Pt(12)
        r.font.color.rgb = WHITE if txt_col != 'FFFFFF' else DARK_BLUE
        sp = c.add_paragraph(content); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.runs[0].font.size = Pt(10)
    else:
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(content).font.size = Pt(10)
figure_caption(doc, "Figure 1 : Architecture globale du systeme Smart Chantier (3-Tiers)")

section_title(doc, "2.4.2 Architecture Flutter (Pattern Provider)", 2)
p = doc.add_paragraph(
    "L'application Flutter adopte le pattern Provider (architecture MVVM simplifiee) pour la "
    "gestion d'etat. Cette approche separe clairement la logique metier de la couche d'affichage :"
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=6)
p.runs[0].font.size = Pt(12)
for layer, items in [
    ("Couche UI (Screens & Widgets)", [
        "DashboardScreen : tableau de bord principal avec statistiques et projets en cours",
        "ProjectsScreen / ProjectDetailScreen : liste et detail des projets",
        "TasksScreen, ReportsScreen, ExpensesScreen, AlertsScreen, ProfileScreen",
        "Widgets reutilisables : StatusBadge, AlertCard, ProjectMiniCard",
    ]),
    ("Couche State Management (Providers)", [
        "AuthProvider : gestion de l'authentification, auto-login, profil JWT",
        "DataProvider : donnees partagees (projets, taches, rapports, depenses, alertes)",
    ]),
    ("Couche Services (API)", [
        "ApiService : client HTTP Dio, toutes les requetes vers le backend Django",
        "SecureStorage : stockage chiffre des tokens JWT",
    ]),
    ("Couche Modeles", [
        "Project, Task, DailyReport, Expense, Alert, SitePhoto, AppUser",
        "Methode fromJson() pour la deserialisation depuis l'API REST",
    ]),
]:
    p = doc.add_paragraph()
    set_para_spacing(p, before=6, after=2)
    r = p.add_run(layer); r.bold = True; r.font.size = Pt(11); r.font.color.rgb = DARK_BLUE
    for item in items:
        bullet_item(doc, item, 1.5)

section_title(doc, "2.4.3 Architecture Django REST Framework", 2)
p = doc.add_paragraph(
    "Le backend suit l'architecture MVT (Model-View-Template) de Django, adaptee pour "
    "une API REST avec DRF. Chaque fonctionnalite est encapsulee dans un module independant (app) :"
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=6)
p.runs[0].font.size = Pt(12)
make_table(doc, ["Module (apps/)", "Responsabilite"], [
    ("accounts",  "Gestion des utilisateurs, authentification JWT, profils et organisations"),
    ("projects",  "CRUD projets, calcul budget, statistiques, suivi avancement"),
    ("tasks",     "Gestion des taches, assignation, suivi statut et priorite"),
    ("reports",   "Rapports journaliers, validation, archivage"),
    ("expenses",  "Depenses, categorisation, controle budgetaire"),
    ("alerts",    "Creation, envoi, resolution et archivage des alertes"),
    ("photos",    "Upload, stockage et service des photos de chantier"),
], col_widths=[4, 12])
figure_caption(doc, "Figure 7 : Architecture Django REST Framework - Structure modulaire")

section_title(doc, "2.4.4 Modele de base de donnees (ERD)", 2)
p = doc.add_paragraph(
    "Le schema de base de donnees de Smart Chantier contient 8 tables principales reliees par "
    "des cles etrangeres. Le tableau suivant presente les entites et leurs relations :"
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=6)
p.runs[0].font.size = Pt(12)
make_table(doc, ["Table", "Cle Primaire", "Cles Etrangeres"], [
    ("accounts_user",         "id (PK)", "organization_id -> projects_organization"),
    ("projects_organization", "id (PK)", "--"),
    ("projects_project",      "id (PK)", "organization_id, created_by_id -> user"),
    ("tasks_task",            "id (PK)", "project_id, assigned_to_id -> user"),
    ("reports_dailyreport",   "id (PK)", "project_id, reporter_id -> user"),
    ("expenses_expense",      "id (PK)", "project_id, approved_by_id -> user"),
    ("alerts_alert",          "id (PK)", "project_id, task_id, created_by_id -> user"),
    ("photos_sitephoto",      "id (PK)", "project_id, report_id, uploaded_by_id -> user"),
], col_widths=[5, 3, 8])
figure_caption(doc, "Figure 5 : Modele Entite-Association (ERD) - Smart Chantier")

section_title(doc, "2.5 Conclusion", 1)
conclusion_intro(doc,
    "L'analyse des besoins a permis d'identifier les principales fonctionnalites attendues "
    "par les differents acteurs du systeme ainsi que les contraintes non fonctionnelles "
    "necessaires a une bonne experience d'utilisation. Grace a la modelisation UML "
    "(diagramme de cas d'utilisation, diagramme de classes, diagramme de sequence) et a "
    "la definition de l'architecture en couches, les interactions avec le systeme ont ete "
    "clairement definies, tout comme la structure des donnees et les flux d'information. "
    "Cette phase preparatoire constitue une base solide pour assurer un developpement "
    "coherent, efficace et adapte aux objectifs du projet."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPITRE 3 : REALISATION ET IMPLEMENTATION
# ═══════════════════════════════════════════════════════════════════════════════
chapter_title(doc, "Chapitre 3 : Realisation et Implementation")

section_title(doc, "3.1 Introduction", 1)
p = doc.add_paragraph(
    "Ce chapitre presente l'environnement technique utilise pour le developpement de "
    "Smart Chantier, les differentes interfaces de l'application, les resultats des tests "
    "realises, ainsi que les apports de ce projet en termes de competences acquises. "
    "Il constitue la partie la plus concrète du rapport, illustrant le travail effectif "
    "de developpement et de mise en oeuvre de la solution."
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=10)
p.runs[0].font.size = Pt(12)

# ── 3.2 Environnement de travail ──
section_title(doc, "3.2 Environnement de travail", 1)
section_title(doc, "3.2.1 Outils de conception", 2)
p = doc.add_paragraph(
    "Dans cette section, nous presentons les outils utilises pour la conception et la modelisation du projet."
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=8)
p.runs[0].font.size = Pt(12)

for outil, logo_fig, desc in [
    ("UML (Unified Modeling Language)", "Figure 8 : Logo UML",
     "Nous avons utilise le langage UML pour modeliser l'application de maniere structuree "
     "et orientee objet. Quatre types de diagrammes ont ete realises : cas d'utilisation, "
     "sequence, classes et modele entite-association. Chacun permet de visualiser une dimension "
     "differente du systeme. Cette approche a facilite la conception et la comprehension globale "
     "de l'application avant son developpement."),
    ("Draw.io", "Figure 9 : Logo Draw.io",
     "Draw.io est un logiciel de dessin de diagrammes en ligne qui offre une interface "
     "conviviale pour creer des diagrammes bases sur differents langages de modelisation, "
     "y compris UML. Il propose une large gamme de formes, de symboles et d'outils de dessin "
     "pour representer les elements du systeme et creer des diagrammes UML de maniere visuelle. "
     "Nous l'avons utilise pour realiser tous les diagrammes de ce rapport."),
]:
    section_title(doc, outil, 3)
    # Placeholder logo
    t_logo = doc.add_table(rows=1, cols=1)
    t_logo.alignment = WD_TABLE_ALIGNMENT.RIGHT
    set_table_borders(t_logo, 'CCCCCC', 4)
    c = t_logo.rows[0].cells[0]
    set_cell_bg(c, 'F5F5F5'); cell_margins(c, 100, 100, 100, 100)
    p2 = c.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"[Logo]\n{outil}"); r2.font.size = Pt(9); r2.italic = True
    figure_caption(doc, logo_fig)
    p = doc.add_paragraph(desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=10)
    p.runs[0].font.size = Pt(12)

section_title(doc, "3.2.2 Outils de developpement", 2)
p = doc.add_paragraph(
    "Dans cette section, nous presentons les principaux langages, frameworks et outils "
    "qui ont ete essentiels au developpement de Smart Chantier."
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=8)
p.runs[0].font.size = Pt(12)

for tech, fig_lbl, desc in [
    ("Langage Dart",
     "Figure 11 : Logo Dart",
     "Le projet a ete developpe en utilisant Dart, un langage oriente objet concu par Google. "
     "Il a servi a ecrire toute la logique de l'interface utilisateur dans Flutter, permettant "
     "un developpement fluide, structure et performant de l'application."),
    ("Framework Flutter",
     "Figure 10 : Logo Flutter",
     "Flutter est un framework open source de Google utilise pour creer des interfaces mobiles "
     "et web modernes. Il permet un developpement rapide avec un seul code source en utilisant "
     "le langage Dart. Il offre d'excellentes performances grace a son moteur natif et au hot reload. "
     "Nous l'avons utilise pour developper toute l'interface utilisateur de Smart Chantier, "
     "deployable sur Android, iOS, Chrome et Edge."),
    ("Django REST Framework",
     "Figure 13 : Logo Django",
     "Django REST Framework (DRF) est une boite a outils puissante pour construire des API Web "
     "avec Django (Python). Nous l'avons utilise pour developper le backend de Smart Chantier, "
     "incluant l'authentification JWT, la gestion des permissions, la serialisation des donnees "
     "et la documentation des endpoints API."),
    ("Visual Studio Code",
     "Figure 12 : Logo Visual Studio Code",
     "Visual Studio Code est un editeur de code source moderne developpe par Microsoft. "
     "Il est leger, rapide et extensible. Il offre des fonctionnalites avancees telles que "
     "la coloration syntaxique, l'auto-completion, le debogage integre, et l'integration avec Git. "
     "Nous l'avons utilise pour ecrire, organiser et maintenir le code Dart et Python du projet."),
]:
    section_title(doc, tech, 3)
    t_logo = doc.add_table(rows=1, cols=1)
    t_logo.alignment = WD_TABLE_ALIGNMENT.RIGHT
    set_table_borders(t_logo, 'CCCCCC', 4)
    c = t_logo.rows[0].cells[0]
    set_cell_bg(c, 'F5F5F5'); cell_margins(c, 80, 80, 80, 80)
    p2 = c.paragraphs[0]; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"[Logo]\n{tech}"); r2.font.size = Pt(8); r2.italic = True
    figure_caption(doc, fig_lbl)
    p = doc.add_paragraph(desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=10)
    p.runs[0].font.size = Pt(12)

table_caption(doc, "Tableau 4 : Technologies et outils utilises")
make_table(doc, ["Technologie / Outil", "Version", "Role"], [
    ("Flutter / Dart",                "3.x / 3.x",  "Framework front-end multi-plateforme"),
    ("Django",                        "4.2 LTS",     "Framework back-end Python"),
    ("Django REST Framework",         "3.14+",       "Construction de l'API REST"),
    ("djangorestframework-simplejwt", "5.x",         "Authentification par tokens JWT"),
    ("Dio",                           "5.x",         "Client HTTP pour Flutter"),
    ("flutter_secure_storage",        "9.x",         "Stockage chiffre des tokens"),
    ("Provider",                      "6.x",         "Gestion d'etat Flutter (MVVM)"),
    ("SQLite / PostgreSQL",           "--",          "Base de donnees dev / production"),
    ("Visual Studio Code",            "--",          "IDE de developpement principal"),
    ("Postman",                       "--",          "Test des endpoints API"),
    ("Git",                           "--",          "Controle de version"),
], col_widths=[5.5, 3, 7.5])

# ── 3.3 Interfaces ──
section_title(doc, "3.3 Presentation des interfaces", 1)
p = doc.add_paragraph(
    "L'interface utilisateur de Smart Chantier a ete decoupee en plusieurs ecrans principaux, "
    "chacun dedie a une fonctionnalite specifique de l'application."
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=10)
p.runs[0].font.size = Pt(12)

def interface(doc, num, title, fig_nums, desc, features, endpoints=None):
    section_title(doc, f"3.3.{num} {title}", 2)
    p = doc.add_paragraph(desc)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=6)
    p.runs[0].font.size = Pt(12)
    for feat in features:
        bullet_item(doc, feat)
    if endpoints:
        p2 = doc.add_paragraph()
        r2 = p2.add_run("Endpoints API utilises :"); r2.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = DARK_BLUE
        for ep in endpoints:
            p3 = doc.add_paragraph(f"    {ep}")
            p3.runs[0].font.name = 'Courier New'; p3.runs[0].font.size = Pt(10)
            set_para_spacing(p3, before=0, after=2)
    for fig_num, fig_title in fig_nums:
        screenshot_box(doc, f"{title} - {fig_title}", fig_num)

interface(doc, 1, "Interface de connexion", [("Figure 14", "Ecran de connexion")],
    "L'ecran de connexion est le point d'entree securise de l'application Smart Chantier. "
    "Il permet aux utilisateurs de s'authentifier avec leur adresse email et mot de passe. "
    "En arriere-plan, l'AuthProvider tente une connexion automatique si un token valide est "
    "deja stocke. En cas d'echec (serveur indisponible), un profil administrateur par defaut "
    "permet l'acces immediat a l'application.",
    [
        "Formulaire de connexion avec validation des champs email et mot de passe.",
        "Affichage / masquage du mot de passe avec icone dediee.",
        "Authentification JWT : envoi des credentials, reception des tokens.",
        "Auto-login transparent si un token valide est deja stocke.",
        "Mode hors-ligne : acces avec donnees par defaut si le serveur est indisponible.",
        "Redirection automatique vers le tableau de bord apres connexion.",
    ],
    ["POST /api/auth/login/ --> {access, refresh, user_data}"]
)

interface(doc, 2, "Tableau de bord", [("Figure 15", "Tableau de bord principal")],
    "Le tableau de bord est l'ecran principal offrant une vue synthetique et en temps reel "
    "de l'etat de tous les projets et activites. La conception est compacte pour maximiser "
    "l'information visible d'un seul coup d'oeil, avec un acces rapide a toutes les fonctions.",
    [
        "AppBar compact avec salutation personnalisee et date du jour en francais.",
        "Carte budget : budget total gere, depenses, pourcentage utilise, barre de progression.",
        "4 indicateurs cles : Projets actifs, Taches en cours, Alertes non lues, Rapports.",
        "Liste des projets en cours avec mini-carte (nom, localisation, budget, progression).",
        "Alertes recentes non resolues avec niveau de criticite colore.",
        "Pull-to-refresh pour actualiser les donnees depuis l'API Django.",
    ],
    ["GET /api/dashboard/", "GET /api/projects/", "GET /api/alerts/"]
)

interface(doc, 3, "Gestion des projets",
    [("Figure 16", "Liste des projets"), ("Figure 17", "Detail d'un projet")],
    "L'ecran des projets affiche la liste complete de tous les projets de l'organisation. "
    "Chaque projet est presente sous forme de carte avec les informations essentielles. "
    "L'ecran de detail centralise toutes les informations relatives a un projet avec des "
    "onglets pour les taches, rapports, depenses et photos associees.",
    [
        "Liste des projets avec filtrage par statut.",
        "Carte projet : nom, localisation, budget, dates de debut/fin, pourcentage d'avancement.",
        "Badge de statut colore : En attente (gris), En cours (bleu), Termine (vert), Annule (rouge).",
        "Ecran de detail avec onglets : Vue generale, Taches, Rapports, Depenses.",
        "Calcul en temps reel du budget restant et detection des depassements.",
        "Navigation directe depuis le tableau de bord par pression sur la carte projet.",
    ],
    ["GET /api/projects/", "GET /api/projects/{id}/", "GET /api/projects/{id}/statistics/"]
)

interface(doc, 4, "Gestion des taches", [("Figure 18", "Liste des taches")],
    "L'ecran des taches permet de suivre l'avancement de chaque activite sur les chantiers. "
    "Les taches sont organisees par projet et classees par priorite avec des indicateurs visuels.",
    [
        "Liste des taches avec filtrage par projet, statut et priorite.",
        "Indicateurs de priorite colores : Faible (vert), Normale (bleu), Haute (orange), Urgente (rouge).",
        "Barre de progression pour chaque tache (0-100%).",
        "Date d'echeance avec alerte visuelle si depassee.",
        "Mise a jour du statut et du pourcentage directement depuis la liste.",
        "Affichage du responsable assigne a chaque tache.",
    ],
    ["GET /api/tasks/?project={id}", "PATCH /api/tasks/{id}/"]
)

interface(doc, 5, "Rapports journaliers", [("Figure 19", "Liste des rapports")],
    "L'ecran des rapports journaliers permet de documenter l'avancement quotidien du chantier. "
    "Chaque rapport enregistre les activites de la journee, les conditions meteo, "
    "le nombre d'ouvriers presents et les incidents eventuels.",
    [
        "Liste des rapports classee par date (plus recent en premier).",
        "Formulaire de creation : date, meteo, effectif, travaux realises, problemes rencontres.",
        "Attachement de photos au rapport de chantier.",
        "Filtrage par projet et par periode.",
    ],
    ["GET /api/reports/?project={id}", "POST /api/reports/"]
)

interface(doc, 6, "Gestion des depenses", [("Figure 20", "Liste des depenses")],
    "L'ecran de gestion des depenses permet de saisir et suivre toutes les sorties budgetaires "
    "liees a un projet. Il offre une visibilite complete sur l'utilisation du budget alloue.",
    [
        "Liste des depenses avec categorie, montant et date.",
        "Categories : Materiaux, Main d'oeuvre, Equipement, Transport, Sous-traitance, Divers.",
        "Formulaire d'ajout avec selection du projet (liste deroulante).",
        "Indicateur de depassement budgetaire avec alerte visuelle.",
        "Filtrage par projet, categorie et periode.",
    ],
    ["GET /api/expenses/?project={id}", "POST /api/expenses/"]
)

interface(doc, 7, "Systeme d'alertes", [("Figure 21", "Liste des alertes")],
    "Le systeme d'alertes centralise toutes les notifications importantes concernant les projets. "
    "Les alertes peuvent etre creees automatiquement (depassement budget, retard) ou "
    "manuellement par les utilisateurs.",
    [
        "Liste des alertes avec niveau de criticite colore.",
        "Niveaux : Information (bleu), Avertissement (orange), Danger (rouge), Critique (violet).",
        "Types : Budget, Retard, Qualite, Securite, Meteo, Administratif.",
        "Bouton de resolution des alertes actives.",
        "Badge counter dans la navigation (nombre d'alertes non lues).",
        "Historique complet des alertes resolues.",
    ],
    ["GET /api/alerts/", "POST /api/alerts/{id}/resolve/"]
)

# ── 3.4 Tests ──
section_title(doc, "3.4 Tests et validation", 1)
p = doc.add_paragraph(
    "Deux types de tests ont ete realises : des tests des API avec Postman et des tests "
    "fonctionnels de l'interface sur navigateur web (Microsoft Edge) et emulateur Android."
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=8)
p.runs[0].font.size = Pt(12)

section_title(doc, "Tests API avec Postman", 3)
table_caption(doc, "Tableau 5 : Resultats des tests API")
make_table(doc, ["Endpoint", "Methode", "Resultat", "Statut"], [
    ("/api/auth/login/",          "POST",   "Connexion, retourne access + refresh tokens + donnees user", "OK"),
    ("/api/projects/",            "GET",    "Liste des projets avec tous les champs attendus",              "OK"),
    ("/api/tasks/",               "GET",    "Liste des taches avec filtrage par projet",                    "OK"),
    ("/api/reports/",             "POST",   "Creation rapport avec authentification Bearer",               "OK"),
    ("/api/expenses/",            "POST",   "Ajout depense avec validation du budget",                      "OK"),
    ("/api/alerts/",              "GET",    "Liste alertes actives et resolues",                            "OK"),
    ("/api/alerts/{id}/resolve/", "POST",   "Resolution d'une alerte active",                              "OK"),
    ("/api/dashboard/",           "GET",    "Statistiques globales pour le tableau de bord",               "OK"),
    ("/api/accounts/me/",         "GET",    "Profil de l'utilisateur connecte",                            "OK"),
    ("/api/auth/refresh/",        "POST",   "Renouvellement du token d'acces via refresh token",           "OK"),
], col_widths=[5.5, 2, 7, 1.5])

section_title(doc, "Tests fonctionnels de l'application", 3)
p = doc.add_paragraph(
    "Les tests fonctionnels ont ete effectues sur navigateur Microsoft Edge et sur emulateur "
    "Android pour valider le comportement de chaque fonctionnalite :"
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=6)
p.runs[0].font.size = Pt(12)
make_table(doc, ["ID", "Cas de test", "Action realise", "Statut"], [
    ("TC01", "Connexion administrateur",   "Saisir email/mdp valides -> Redirection dashboard",   "PASS"),
    ("TC02", "Auto-login",                 "Relancer app avec token stocke -> Connexion directe", "PASS"),
    ("TC03", "Mode hors-ligne",            "Couper serveur -> Donnees par defaut affichees",       "PASS"),
    ("TC04", "Tableau de bord",            "Ouvrir app -> Budget, stats, projets, alertes",        "PASS"),
    ("TC05", "Navigation vers projet",     "Tapper un projet -> Ecran detail avec onglets",        "PASS"),
    ("TC06", "Creer un rapport",           "Rapports -> Nouveau -> Remplir -> Rapport ajoute",     "PASS"),
    ("TC07", "Ajouter une depense",        "Depenses -> Ajouter -> Budget mis a jour",             "PASS"),
    ("TC08", "Resoudre une alerte",        "Alertes -> Resoudre -> Statut = Resolu",               "PASS"),
    ("TC09", "Pull-to-refresh",            "Glisser bas -> Rechargement depuis API",               "PASS"),
    ("TC10", "Deconnexion",                "Profil -> Deconnecter -> Retour ecran connexion",      "PASS"),
], col_widths=[1.5, 4, 7, 1.5])

# ── 3.5 Apports ──
section_title(doc, "3.5 Les apports du projet", 1)
p = doc.add_paragraph(
    "Ce projet de fin d'etudes a ete pour nous une opportunite precieuse d'acquerir et de "
    "renforcer un ensemble de competences essentielles."
)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=6)
p.runs[0].font.size = Pt(12)

section_title(doc, "Competences techniques developpees :", 3)
for item in [
    "Maitrise de Flutter et Dart pour le developpement d'interfaces mobiles et web performantes.",
    "Conception et developpement d'une API REST robuste avec Django REST Framework.",
    "Implementation d'un systeme d'authentification securise par tokens JWT.",
    "Modelisation UML pour structurer les projets et formaliser les besoins.",
    "Gestion d'une base de donnees relationnelle avec l'ORM Django.",
    "Integration de methodes de gestion de projet (Git, documentation, tests).",
]:
    bullet_item(doc, item)

section_title(doc, "Competences transversales acquises :", 3)
for item in [
    "Travail en autonomie, gestion du temps et respect des echeances.",
    "Capacite a analyser les problemes rencontres et a proposer des solutions adaptees.",
    "Amelioration des competences en documentation et redaction technique.",
    "Curiosite naturelle stimulee par la decouverte de nouvelles technologies.",
]:
    bullet_item(doc, item)

section_title(doc, "3.6 Conclusion", 1)
conclusion_intro(doc,
    "Ce chapitre nous a permis de presenter l'environnement de developpement utilise, "
    "les differentes interfaces realisees pour Smart Chantier, et les resultats des tests "
    "effectues. L'ensemble des fonctionnalites definies dans l'analyse des besoins a ete "
    "correctement implemente et valide. Les tests fonctionnels confirment que l'application "
    "repond aux attentes des utilisateurs et fonctionne correctement sur les differentes "
    "plateformes cibles. Ce projet a egalement ete l'occasion d'acquerir des competences "
    "techniques et transversales importantes pour notre parcours professionnel."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CONCLUSION GENERALE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p, before=0, after=20)
r = p.add_run("Conclusion Generale")
r.font.size = Pt(18); r.bold = True; r.font.color.rgb = DARK_BLUE
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot = OxmlElement('w:bottom'); bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '12')
bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), '003366')
pBdr.append(bot); pPr.append(pBdr)
doc.add_paragraph()

for para in [
    "Ce projet de fin d'etudes, realise dans le cadre de notre formation en Genie Logiciel, "
    "avait pour principal objectif la conception et le developpement d'une application mobile "
    "et web dediee a la gestion intelligente des chantiers de construction : Smart Chantier.",

    "Durant ce projet, nous avons participe a toutes les etapes : analyse des besoins, "
    "conception de l'architecture, modelisation UML, developpement de l'interface avec Flutter, "
    "developpement du back-end avec Django REST Framework, integration des principales "
    "fonctionnalites, ainsi que realisation de tests pour garantir une bonne experience utilisateur.",

    "Sur le plan technique, ce projet nous a permis de maitriser des outils concrets du "
    "developpement mobile et web, d'approfondir nos competences en conception logicielle "
    "et en architecture de systemes distribues. Sur le plan personnel, nous avons developpe "
    "davantage d'autonomie, de reactivite, et une meilleure capacite d'adaptation.",

    "Quelques difficultes ont ete rencontrees, notamment lors de la gestion des tokens JWT "
    "entre Flutter et Django, la resolution des problemes de CORS, ou encore l'adaptation "
    "de l'interface pour le web et le mobile avec un seul code source. Ces defis ont cependant "
    "ete surmontes et ont represente autant d'opportunites d'apprentissage.",
]:
    p = doc.add_paragraph(para)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; set_para_spacing(p, after=10)
    p.runs[0].font.size = Pt(12)

p = doc.add_paragraph()
r = p.add_run("Perspectives"); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = DARK_BLUE
set_para_spacing(p, before=10, after=6)

for item in [
    "Connexion a un serveur cloud (AWS, Azure) avec base de donnees PostgreSQL.",
    "Notifications push en temps reel via Firebase Cloud Messaging.",
    "Module de planification avec diagramme de Gantt interactif.",
    "Intelligence artificielle pour la prediction des retards et depassements budgetaires.",
    "Developpement d'une interface administrateur web avancee.",
    "Deploiement sur Google Play Store et Apple App Store.",
    "Support multilingue : Arabe, Francais, Anglais.",
]:
    bullet_item(doc, item)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# BIBLIOGRAPHIE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p, before=0, after=20)
r = p.add_run("Bibliographie")
r.font.size = Pt(18); r.bold = True; r.font.color.rgb = DARK_BLUE
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot = OxmlElement('w:bottom'); bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '12')
bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), '003366')
pBdr.append(bot); pPr.append(pBdr)
doc.add_paragraph()

for ref, title, source in [
    ("[1]", "Flutter - Documentation officielle",              "https://flutter.dev/docs"),
    ("[2]", "Dart - Documentation officielle",                 "https://dart.dev/guides"),
    ("[3]", "Django REST Framework",                           "https://www.django-rest-framework.org"),
    ("[4]", "djangorestframework-simplejwt",                   "https://django-rest-framework-simplejwt.readthedocs.io"),
    ("[5]", "Package Dio - Client HTTP pour Flutter",          "https://pub.dev/packages/dio"),
    ("[6]", "Package flutter_secure_storage",                  "https://pub.dev/packages/flutter_secure_storage"),
    ("[7]", "Package Provider - Gestion d'etat Flutter",       "https://pub.dev/packages/provider"),
    ("[8]", "Django Documentation 4.2 LTS",                    "https://docs.djangoproject.com/en/4.2/"),
    ("[9]", "JSON Web Tokens - RFC 7519",                      "https://tools.ietf.org/html/rfc7519"),
    ("[10]","OWASP - API Security Top 10",                     "https://owasp.org/www-project-api-security/"),
    ("[11]","Visual Studio Code",                              "https://code.visualstudio.com"),
    ("[12]","Android Studio",                                  "https://developer.android.com/studio"),
    ("[13]","Robert C. Martin - Clean Architecture",           "Prentice Hall, 2017"),
    ("[14]","Cours et tutoriels Flutter - Flutter API Docs",   "https://api.flutter.dev/index.html"),
]:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{ref} "); r1.bold = True; r1.font.size = Pt(12); r1.font.color.rgb = DARK_BLUE
    r2 = p.add_run(f"{title}. ")
    r2.font.size = Pt(12)
    r3 = p.add_run(source); r3.font.size = Pt(11); r3.font.color.rgb = RGBColor(0x00, 0x56, 0xA0); r3.underline = True
    set_para_spacing(p, before=0, after=4)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ANNEXES
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p, before=0, after=20)
r = p.add_run("Annexes")
r.font.size = Pt(18); r.bold = True; r.font.color.rgb = DARK_BLUE
pPr = p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bot = OxmlElement('w:bottom'); bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '12')
bot.set(qn('w:space'), '4'); bot.set(qn('w:color'), '003366')
pBdr.append(bot); pPr.append(pBdr)
doc.add_paragraph()

section_title(doc, "Annexe A : Structure du projet Flutter", 2)
p = doc.add_paragraph()
r = p.add_run(
"""flutter_app/lib/
  main.dart                    # Point d'entree de l'application
  core/
    constants.dart             # URL API, constantes globales
    default_data.dart          # Donnees par defaut (mode hors-ligne)
    theme.dart                 # Couleurs et styles de l'application
  models/
    project.dart  task.dart  daily_report.dart
    expense.dart  alert.dart  site_photo.dart  app_user.dart
  providers/
    auth_provider.dart         # Authentification, auto-login, JWT
    data_provider.dart         # Donnees partagees et chargement API
  services/
    api_service.dart           # Toutes les requetes HTTP (Dio)
  screens/
    dashboard/  projects/  tasks/  reports/
    expenses/  alerts/  profile/
    main_screen.dart           # Navigation principale (BottomNavigationBar)
  widgets/
    alert_card.dart  status_badge.dart
pubspec.yaml                   # Dependances du projet"""
)
r.font.name = 'Courier New'; r.font.size = Pt(9)

doc.add_paragraph()
section_title(doc, "Annexe B : Structure du backend Django", 2)
p = doc.add_paragraph()
r = p.add_run(
"""django_backend/
  config/
    settings.py      # Configuration (DB, CORS, JWT, REST_FRAMEWORK)
    urls.py          # Routage principal des endpoints
  apps/
    accounts/        # Modele User personnalise, auth JWT
    projects/        # Projets et organisations
    tasks/           # Taches et affectations
    reports/         # Rapports journaliers
    expenses/        # Depenses et categories
    alerts/          # Alertes et notifications
    photos/          # Photos de chantier
  manage.py
  seed_data.py       # Script pour remplir la base avec des donnees d'exemple
  requirements.txt   # Dependances Python (Django, DRF, simplejwt, ...)"""
)
r.font.name = 'Courier New'; r.font.size = Pt(9)

doc.add_paragraph()
section_title(doc, "Annexe C : Endpoints API complets", 2)
table_caption(doc, "Tableau 6 : Endpoints API - Liste complete")
make_table(doc, ["Methode", "Endpoint", "Description"], [
    ("POST",   "/api/auth/login/",           "Connexion, retourne JWT tokens + donnees utilisateur"),
    ("POST",   "/api/auth/refresh/",         "Renouvellement de l'access token"),
    ("POST",   "/api/auth/logout/",          "Deconnexion, invalidation du token"),
    ("GET",    "/api/accounts/me/",          "Profil de l'utilisateur connecte"),
    ("PUT",    "/api/accounts/me/",          "Modification du profil utilisateur"),
    ("GET",    "/api/projects/",             "Liste tous les projets de l'organisation"),
    ("POST",   "/api/projects/",             "Creer un nouveau projet"),
    ("GET",    "/api/projects/{id}/",        "Detail d'un projet specifique"),
    ("PUT",    "/api/projects/{id}/",        "Modifier un projet"),
    ("DELETE", "/api/projects/{id}/",        "Supprimer un projet"),
    ("GET",    "/api/tasks/",                "Liste des taches (filtre: ?project=id)"),
    ("POST",   "/api/tasks/",                "Creer une nouvelle tache"),
    ("PATCH",  "/api/tasks/{id}/",           "Mise a jour partielle (statut, pourcentage)"),
    ("GET",    "/api/reports/",              "Liste des rapports (filtre: ?project=id)"),
    ("POST",   "/api/reports/",              "Creer un rapport journalier"),
    ("GET",    "/api/expenses/",             "Liste des depenses (filtre: ?project=id)"),
    ("POST",   "/api/expenses/",             "Enregistrer une nouvelle depense"),
    ("GET",    "/api/alerts/",               "Liste des alertes (filtre: ?status=active)"),
    ("POST",   "/api/alerts/{id}/resolve/",  "Resoudre une alerte active"),
    ("GET",    "/api/dashboard/",            "Statistiques globales pour le tableau de bord"),
    ("GET",    "/api/photos/",               "Liste des photos de chantier"),
    ("POST",   "/api/photos/",               "Uploader une photo de chantier"),
], col_widths=[2, 5.5, 8.5])

# ── Sauvegarde ──────────────────────────────────────────────────────────────
out = r"C:\Users\INDEX INFORMATIQUE\Downloads\PFE\Rapport_PFE_Smart_Chantier.docx"
doc.save(out)
print("Rapport genere avec succes : " + out)
print("Pages estimees : ~60 pages")
print("Structure : Page de garde, Dedicace, Remerciements, TDM, Liste figures,")
print("            Liste tableaux, Abreviations, Introduction,")
print("            Chapitre 1 : Presentation du Projet,")
print("            Chapitre 2 : Analyse des Besoins et Conception,")
print("            Chapitre 3 : Realisation et Implementation,")
print("            Conclusion Generale, Bibliographie, Annexes A-C")
